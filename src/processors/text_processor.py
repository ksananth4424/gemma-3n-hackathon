"""
Text Content Processor for various text formats
"""

import logging
import docx
from pathlib import Path
from typing import Optional
from striprtf.striprtf import rtf_to_text

from ..utils.config_manager import ConfigManager

logger = logging.getLogger('accessibility_assistant.text_processor')


class TextProcessor:
    """
    Text content processor for various text formats
    """
    
    def __init__(self, config_manager: ConfigManager):
        self.config = config_manager
        
        # Supported encodings to try
        self.encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1']
        
    def extract_content(self, file_path: str) -> str:
        """
        Extract content from text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            Extracted text content
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.docx':
                return self._extract_docx_content(file_path)
            elif file_extension == '.rtf':
                return self._extract_rtf_content(file_path)
            else:
                return self._extract_plain_text(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text content from {file_path}: {e}")
            raise
    
    def _extract_plain_text(self, file_path: str) -> str:
        """
        Extract content from plain text files with encoding detection
        
        Args:
            file_path: Path to text file
            
        Returns:
            Text content
        """
        content = None
        
        # Try different encodings
        for encoding in self.encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.debug(f"Successfully read {file_path} with {encoding} encoding")
                break
                
            except UnicodeDecodeError:
                logger.debug(f"Failed to read {file_path} with {encoding} encoding")
                continue
            except Exception as e:
                logger.error(f"Error reading {file_path} with {encoding}: {e}")
                continue
        
        if content is None:
            # Last resort: read as binary and decode with errors='ignore'
            try:
                with open(file_path, 'rb') as f:
                    raw_content = f.read()
                content = raw_content.decode('utf-8', errors='ignore')
                logger.warning(f"Used fallback decoding for {file_path}")
            except Exception as e:
                logger.error(f"Fallback decoding failed for {file_path}: {e}")
                raise
        
        return self._clean_text_content(content)
    
    def _extract_docx_content(self, file_path: str) -> str:
        """
        Extract content from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Text content
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            
            content = '\n'.join(paragraphs)
            logger.info(f"Extracted {len(content)} characters from DOCX")
            
            return self._clean_text_content(content)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            raise
    
    def _extract_rtf_content(self, file_path: str) -> str:
        """
        Extract content from RTF file
        
        Args:
            file_path: Path to RTF file
            
        Returns:
            Text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            # Convert RTF to plain text
            text_content = rtf_to_text(rtf_content)
            
            logger.info(f"Extracted {len(text_content)} characters from RTF")
            
            return self._clean_text_content(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting RTF content: {e}")
            raise
    
    def _clean_text_content(self, content: str) -> str:
        """
        Clean and normalize text content
        
        Args:
            content: Raw text content
            
        Returns:
            Cleaned text
        """
        if not content:
            return ""
        
        # Normalize line endings
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace while preserving structure
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove trailing/leading whitespace
            cleaned_line = line.strip()
            
            # Skip empty lines but preserve paragraph breaks
            if cleaned_line or (cleaned_lines and cleaned_lines[-1]):
                cleaned_lines.append(cleaned_line)
        
        # Join lines and remove excessive blank lines
        cleaned_content = '\n'.join(cleaned_lines)
        
        # Remove more than 2 consecutive newlines
        import re
        cleaned_content = re.sub(r'\n{3,}', '\n\n', cleaned_content)
        
        return cleaned_content.strip()
    
    def get_text_statistics(self, content: str) -> dict:
        """
        Get statistics about the text content
        
        Args:
            content: Text content
            
        Returns:
            Dictionary with text statistics
        """
        if not content:
            return {
                'character_count': 0,
                'word_count': 0,
                'line_count': 0,
                'paragraph_count': 0
            }
        
        # Basic statistics
        character_count = len(content)
        word_count = len(content.split())
        line_count = len(content.split('\n'))
        
        # Count paragraphs (separated by blank lines)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        paragraph_count = len(paragraphs)
        
        return {
            'character_count': character_count,
            'word_count': word_count,
            'line_count': line_count,
            'paragraph_count': paragraph_count,
            'average_words_per_paragraph': word_count / max(paragraph_count, 1)
        }
    
    def estimate_reading_time(self, content: str, wpm: int = 200) -> float:
        """
        Estimate reading time for content
        
        Args:
            content: Text content
            wpm: Words per minute reading speed
            
        Returns:
            Estimated reading time in minutes
        """
        stats = self.get_text_statistics(content)
        word_count = stats['word_count']
        
        reading_time_minutes = word_count / wpm
        return reading_time_minutes
    
    def extract_key_sections(self, content: str) -> dict:
        """
        Extract key sections from structured text
        
        Args:
            content: Text content
            
        Returns:
            Dictionary with extracted sections
        """
        sections = {
            'title': '',
            'headings': [],
            'summary_section': '',
            'conclusion_section': ''
        }
        
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Detect potential title (first significant line)
            if not sections['title'] and len(line) < 100 and i < 5:
                sections['title'] = line
            
            # Detect headings (lines that are short and in caps or title case)
            if (len(line) < 80 and 
                (line.isupper() or line.istitle()) and 
                not line.endswith('.') and
                len(line.split()) <= 8):
                sections['headings'].append(line)
            
            # Look for summary sections
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in ['summary', 'abstract', 'overview']):
                # Get the next few lines as summary
                summary_lines = []
                for j in range(i + 1, min(i + 10, len(lines))):
                    if lines[j].strip():
                        summary_lines.append(lines[j].strip())
                    elif summary_lines:  # Stop at blank line after content
                        break
                sections['summary_section'] = ' '.join(summary_lines)
            
            # Look for conclusion sections
            if any(keyword in line_lower for keyword in ['conclusion', 'summary', 'final']):
                conclusion_lines = []
                for j in range(i + 1, min(i + 10, len(lines))):
                    if lines[j].strip():
                        conclusion_lines.append(lines[j].strip())
                    elif conclusion_lines:
                        break
                sections['conclusion_section'] = ' '.join(conclusion_lines)
        
        return sections
